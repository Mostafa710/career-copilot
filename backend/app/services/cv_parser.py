"""CV parsing and text extraction service supporting PDF, DOCX, pytesseract OCR fallback, and text paste."""

import io
import re
import logging
from typing import Tuple
import pdfplumber
import docx
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


def sanitize_text(text: str) -> str:
    """Clean and normalize extracted resume text."""
    if not text:
        return ""
    # Normalize line breaks and tabs
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Clean leading/trailing spaces per line and multiple inline spaces
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    text = "\n".join(lines)
    # Replace excessive newlines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(content: bytes) -> Tuple[str, bool]:
    """
    Extract text from PDF bytes.
    If extracted text is under 50 characters, triggers pytesseract OCR fallback.
    Returns (extracted_text, used_ocr).
    """
    extracted_text = ""
    used_ocr = False

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_text += page_text + "\n"

        # Check if text is insufficient (indicates scanned / image-based PDF)
        if len(extracted_text.strip()) < 50:
            logger.info("PDF has minimal digital text (<50 chars). Triggering pytesseract OCR fallback.")
            used_ocr = True
            ocr_text = ""
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    try:
                        # Convert page to image
                        page_image = page.to_image(resolution=200).original
                        page_ocr = pytesseract.image_to_string(page_image)
                        ocr_text += page_ocr + "\n"
                    except Exception as ocr_err:
                        logger.warning(f"Error running OCR on PDF page: {ocr_err}")
            
            if len(ocr_text.strip()) > len(extracted_text.strip()):
                extracted_text = ocr_text

    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        # Direct OCR attempt on failure
        try:
            used_ocr = True
            image = Image.open(io.BytesIO(content))
            extracted_text = pytesseract.image_to_string(image)
        except Exception as img_err:
            logger.error(f"Fallback OCR on raw image failed: {img_err}")

    return sanitize_text(extracted_text), used_ocr


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        doc = docx.Document(io.BytesIO(content))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return sanitize_text("\n".join(full_text))
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return ""


def parse_cv_file(filename: str, content: bytes) -> Tuple[str, bool]:
    """
    Parse uploaded CV file based on its extension.
    Returns (extracted_text, used_ocr).
    """
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        return extract_text_from_pdf(content)
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(content), False
    elif ext in ["txt", "md"]:
        return sanitize_text(content.decode("utf-8", errors="ignore")), False
    else:
        # Attempt image OCR if image extension
        if ext in ["png", "jpg", "jpeg", "webp", "tiff"]:
            try:
                img = Image.open(io.BytesIO(content))
                text = pytesseract.image_to_string(img)
                return sanitize_text(text), True
            except Exception as e:
                logger.error(f"Error in image OCR: {e}")
        return "", False
