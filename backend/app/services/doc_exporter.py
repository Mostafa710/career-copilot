"""Document Export Service generating ATS-compliant DOCX and PDF files."""

import io
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List


class DocumentExporter:
    @staticmethod
    def generate_docx_cv(tailored_data: Dict[str, Any], candidate_name: str = "Candidate") -> bytes:
        """Generate a clean, ATS-compliant Word DOCX resume."""
        doc = docx.Document()

        # Set 1-inch margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Header: Name
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(candidate_name)
        name_run.font.name = "Arial"
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(15, 23, 42)  # Slate-900

        # Header: Contact Info
        contact = tailored_data.get("contact_info", {})
        contact_parts = []
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("linkedin_url"):
            contact_parts.append(contact["linkedin_url"])
        if contact.get("github_url"):
            contact_parts.append(contact["github_url"])

        if contact_parts:
            contact_p = doc.add_paragraph()
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_p.add_run(" | ".join(contact_parts))
            contact_run.font.name = "Arial"
            contact_run.font.size = Pt(9.5)
            contact_run.font.color.rgb = RGBColor(71, 85, 105)  # Slate-600

        # Summary Section
        summary = tailored_data.get("professional_summary")
        if summary:
            DocumentExporter._add_section_heading(doc, "PROFESSIONAL SUMMARY")
            sum_p = doc.add_paragraph()
            sum_run = sum_p.add_run(summary)
            sum_run.font.name = "Arial"
            sum_run.font.size = Pt(10)

        # Technical Skills Section
        skills = tailored_data.get("skills", [])
        if skills:
            DocumentExporter._add_section_heading(doc, "TECHNICAL SKILLS")
            skills_p = doc.add_paragraph()
            skills_run = skills_p.add_run(", ".join(skills) if isinstance(skills, list) else str(skills))
            skills_run.font.name = "Arial"
            skills_run.font.size = Pt(10)

        # Work Experience Section
        experiences = tailored_data.get("experience", [])
        if experiences:
            DocumentExporter._add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
            for exp in experiences:
                title = exp.get("title", "")
                company = exp.get("company", "")
                dates = exp.get("dates", "")
                
                job_p = doc.add_paragraph()
                r1 = job_p.add_run(f"{title} — {company}")
                r1.font.name = "Arial"
                r1.font.size = Pt(10.5)
                r1.font.bold = True
                
                if dates:
                    r2 = job_p.add_run(f" ({dates})")
                    r2.font.name = "Arial"
                    r2.font.size = Pt(9.5)
                    r2.font.italic = True
                    r2.font.color.rgb = RGBColor(100, 116, 139)

                bullets = exp.get("bullets", [])
                for b in bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    b_run = bp.add_run(b)
                    b_run.font.name = "Arial"
                    b_run.font.size = Pt(9.5)

        # Education Section
        education = tailored_data.get("education", [])
        if education:
            DocumentExporter._add_section_heading(doc, "EDUCATION")
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("year", "")
                
                edu_p = doc.add_paragraph()
                e_run = edu_p.add_run(f"{degree} — {institution} {f'({year})' if year else ''}")
                e_run.font.name = "Arial"
                e_run.font.size = Pt(10)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _add_section_heading(doc: docx.Document, title: str):
        """Add a standardized ATS section header with divider styling."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59)

    @staticmethod
    def generate_html_cv(tailored_data: Dict[str, Any], candidate_name: str = "Candidate") -> str:
        """Generate a semantic, clean HTML template styled for printing and viewing."""
        contact = tailored_data.get("contact_info", {})
        contact_items = [
            f"<span>{v}</span>" for k, v in contact.items() if v and isinstance(v, str)
        ]
        
        skills_html = ""
        skills = tailored_data.get("skills", [])
        if skills:
            skills_str = ", ".join(skills) if isinstance(skills, list) else str(skills)
            skills_html = f"""
            <section class="section">
                <h2 class="section-title">Technical Skills</h2>
                <p class="skills-list">{skills_str}</p>
            </section>
            """

        experience_html = ""
        experiences = tailored_data.get("experience", [])
        if experiences:
            exp_items = []
            for exp in experiences:
                bullets = "".join([f"<li>{b}</li>" for b in exp.get("bullets", [])])
                exp_items.append(f"""
                <div class="exp-item">
                    <div class="exp-header">
                        <strong>{exp.get('title', '')} — {exp.get('company', '')}</strong>
                        <span class="exp-dates">{exp.get('dates', '')}</span>
                    </div>
                    <ul class="bullet-list">{bullets}</ul>
                </div>
                """)
            experience_html = f"""
            <section class="section">
                <h2 class="section-title">Professional Experience</h2>
                {''.join(exp_items)}
            </section>
            """

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>
    body {{ font-family: 'Helvetica Neue', Arial, sans-serif; line-height: 1.5; color: #1e293b; max-width: 800px; margin: 0 auto; padding: 30px; }}
    .header {{ text-align: center; border-bottom: 2px solid #e2e8f0; padding-bottom: 15px; margin-bottom: 20px; }}
    .name {{ font-size: 24px; font-weight: 700; margin: 0; color: #0f172a; text-transform: uppercase; letter-spacing: 0.5px; }}
    .contact {{ font-size: 13px; color: #64748b; margin-top: 6px; }}
    .contact span:not(:last-child)::after {{ content: " | "; color: #94a3b8; }}
    .section {{ margin-bottom: 20px; }}
    .section-title {{ font-size: 14px; font-weight: 700; text-transform: uppercase; letter-spacing: 1px; color: #334155; border-bottom: 1px solid #cbd5e1; padding-bottom: 4px; margin-bottom: 10px; }}
    .exp-item {{ margin-bottom: 14px; }}
    .exp-header {{ display: flex; justify-content: space-between; font-size: 14px; margin-bottom: 4px; }}
    .exp-dates {{ font-size: 12px; color: #64748b; font-style: italic; }}
    .bullet-list {{ margin: 0; padding-left: 20px; font-size: 13px; color: #334155; }}
    .bullet-list li {{ margin-bottom: 4px; }}
    .skills-list {{ font-size: 13px; color: #334155; margin: 0; }}
</style>
</head>
<body>
    <div class="header">
        <h1 class="name">{candidate_name}</h1>
        <div class="contact">{''.join(contact_items)}</div>
    </div>
    {skills_html}
    {experience_html}
</body>
</html>"""


document_exporter = DocumentExporter()
